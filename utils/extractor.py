import os
import fitz  # PyMuPDF
import docx
import difflib
import logging
from telegram import Update
from telegram.ext import ContextTypes


BASE_DIR = "documentos"
TEMPLATE_DIR = "templates"

# --- Utilidades de lectura ---
def leer_pdf(path):
    text = ""
    try:
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception:
        pass
    return text

def leer_docx(path):
    text = ""
    try:
        doc = docx.Document(path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception:
        pass
    return text

# --- Lectura de tablas específicas ---
def extraer_tablas_docx(path):
    resultados = []
    try:
        doc = docx.Document(path)
        for table in doc.tables:
            encabezados = [cell.text.lower().strip() for cell in table.rows[0].cells]
            if all(col in encabezados for col in ["situación", "incidencias", "procedimiento", "referencia legal 2022"]):
                for row in table.rows[1:]:
                    fila = {encabezados[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
                    resultados.append(fila)
    except Exception:
        pass
    return resultados

# --- Carga inicial de todos los documentos ---
def procesar_documentos(ruta_carpeta):
    base_textos = []
    base_tablas = []

    for root, _, files in os.walk(ruta_carpeta):
        for file in files:
            path = os.path.join(root, file)
            if file.endswith(".pdf"):
                texto = leer_pdf(path)
                base_textos.append((file, texto))
            elif file.endswith(".docx"):
                texto = leer_docx(path)
                tablas = extraer_tablas_docx(path)
                base_textos.append((file, texto))
                base_tablas.extend(tablas)

    return base_textos, base_tablas


# --- Respuesta desde tablas ---
def buscar_en_tablas(pregunta, base_tablas):
    mejores = []
    for fila in base_tablas:
        if isinstance(fila, dict):
            texto = f"{fila.get('situación', '')} {fila.get('incidencias', '')}"
            simil = difflib.SequenceMatcher(None, pregunta.lower(), texto.lower()).ratio()
            if simil > 0.4:
                mejores.append((simil, fila))

    mejores.sort(reverse=True, key=lambda x: x[0])
    if mejores:
        mejor = mejores[0][1]
        return f"Situación: {mejor.get('situación', '')}\nIncidencias: {mejor.get('incidencias', '')}\nProcedimiento: {mejor.get('procedimiento', '')}\nReferencia Legal: {mejor.get('referencia legal 2022', '')}"
    return None

# --- Sugerencia de plantilla ---
def detectar_plantilla(pregunta):
    """
    Detecta si el usuario está solicitando una plantilla o documento, y retorna el archivo si existe.
    """
    palabras_clave = [
        "plantilla", "formato", "documento", "formulario", "modelo", "archivo"
    ]
    if any(palabra in pregunta.lower() for palabra in palabras_clave):
        posibles = []
        for file in os.listdir(TEMPLATE_DIR):
            nombre = file.lower().replace("_", " ").replace("-", " ")
            simil = difflib.SequenceMatcher(None, pregunta.lower(), nombre).ratio()
            if simil > 0.4:
                posibles.append((simil, file))
        posibles.sort(reverse=True)
        return posibles[0][1] if posibles else "NECESITA_SUBIR"
    return None


def buscar_por_situacion(pregunta, base_tablas):
    """
    Busca una coincidencia exacta (o casi exacta) en la columna 'situación'.
    Si encuentra, devuelve una respuesta con el resto de columnas.
    """
    mejores = []
    for fila in base_tablas:
        situacion = fila.get("situación", "").lower().strip()
        simil = difflib.SequenceMatcher(None, pregunta.lower(), situacion).ratio()
        if simil > 0.6:  # puedes ajustar el umbral si es necesario
            mejores.append((simil, fila))

    mejores.sort(reverse=True, key=lambda x: x[0])
    if mejores:
        mejor = mejores[0][1]
        return (
            f"Situación: {mejor['situación']}\n"
            f"Modalidad/Incidencias: {mejor['incidencias']}\n"
            f"Procedimiento: {mejor['procedimiento']}\n"
            f"Referencia Legal: {mejor['referencia legal 2022']}"
        )
    return None
import json

import json

def cargar_json_desde_txt(carpeta):
    datos = []
    for file in os.listdir(carpeta):
        if file.endswith(".txt"):
            path = os.path.join(carpeta, file)
            try:
                with open(path, encoding="utf-8") as f:
                    contenido = json.load(f)
                    if isinstance(contenido, list):
                        datos.extend(contenido)
                    else:
                        datos.append(contenido)
            except Exception as e:
                logging.warning(f"Error cargando {file}: {e}")
    return datos



def buscar_en_json(pregunta, datos_json):
    mejores = []
    for entrada in datos_json:
        texto = f"{entrada.get('situacion', '')} {entrada.get('modalidad', '')}"
        simil = difflib.SequenceMatcher(None, pregunta.lower(), texto.lower()).ratio()
        if simil > 0.4:
            mejores.append((simil, entrada))

    mejores.sort(reverse=True, key=lambda x: x[0])
    if mejores:
        mejor = mejores[0][1]
        return f"Situación: {mejor.get('situacion', '')}\nModalidad: {mejor.get('modalidad', '')}\nProcedimiento: {mejor.get('procedimiento', '')}\nReferencia Legal: {mejor.get('referencia_legal', '')}"
    return None


async def documentar(update: Update, context: ContextTypes.DEFAULT_TYPE):
    mensaje = update.message.text
    tema = mensaje.replace("documentar", "", 1).strip().lower()

    ejemplos = {
        "inspección": [
            "**Acta de Inspección**",
            "Fecha: ____________",
            "Lugar de la inspección: ____________",
            "Unidad actuante: ____________",
            "Funcionario responsable: ____________",
            "Descripción de los hechos: ____________",
            "Resultado de la inspección: ____________",
            "Observaciones: ____________"
        ],
        "detención": [
            "**Acta de Detención**",
            "Fecha y hora: ____________",
            "Lugar de los hechos: ____________",
            "Nombre del detenido: ____________",
            "Cédula de identidad: ____________",
            "Motivo de la detención: ____________",
            "Evidencias incautadas: ____________",
            "Funcionarios actuantes: ____________"
        ],
        "informe": [
            "**Informe Operativo**",
            "Fecha del informe: ____________",
            "Unidad responsable: ____________",
            "Nombre del redactor: ____________",
            "Resumen de la operación: ____________",
            "Resultados obtenidos: ____________",
            "Recomendaciones: ____________"
        ],
        "acta": [
            "**Acta General**",
            "Fecha: ____________",
            "Lugar: ____________",
            "Unidad actuante: ____________",
            "Resumen de los hechos: ____________",
            "Acciones tomadas: ____________",
            "Firmas de los responsables: ____________"
        ]
    }

    # Detectar tipo sugerido
    for clave, campos in ejemplos.items():
        if clave in tema:
            await update.message.reply_text(
                "\n".join(campos),
                parse_mode="Markdown"
            )
            return

    # Si no se encuentra un caso claro, generar una estructura genérica
    estructura_generica = [
        f"**Documento: {tema.title()}**",
        "Fecha: ____________",
        "Lugar: ____________",
        "Unidad o funcionario responsable: ____________",
        "Descripción detallada: ____________",
        "Observaciones: ____________"
    ]
    await update.message.reply_text(
        "\n".join(estructura_generica),
        parse_mode="Markdown"
    )

